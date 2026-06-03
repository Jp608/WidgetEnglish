from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(".artifacts/docs")
USER_DOCX = OUT_DIR / "WidgetEnglish_Guia_usuario_modulos.docx"
ADMIN_DOCX = OUT_DIR / "WidgetEnglish_Estadisticas_admin.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5C6670"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.size = Pt(9)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.text = title
    footer.style = doc.styles["Normal"]
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(section.footer.add_paragraph())


def add_title_block(doc: Document, title: str, subtitle: str, metadata: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(metadata)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(INK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_label_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2700, 6660])
    for row_idx, (label, detail) in enumerate(rows):
        if row_idx > 0:
            table.add_row()
        row = table.rows[row_idx]
        row.cells[0].text = label
        row.cells[1].text = detail
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            row.cells[0].paragraphs[0].runs[0].font.bold = True


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.runs[0].font.bold = True
        paragraph.runs[0].font.size = Pt(10)

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = value
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(10)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title_run.font.size = Pt(10)
    p.add_run("\n")
    body_run = p.add_run(body)
    body_run.font.size = Pt(10)


def build_user_doc() -> None:
    doc = Document()
    style_document(doc, "WidgetEnglish - Guía de usuario y módulos")
    add_title_block(
        doc,
        "WidgetEnglish: guía de usuario y módulos",
        "Descripción funcional de la app, módulos, widget y objetivo diario",
        "Alcance: documento orientado a explicar cómo se usa la app y cómo se comportan sus módulos principales. Base: revisión del código Kotlin/Room/DataStore/Firebase del proyecto.",
    )

    doc.add_heading("1. Descripción general", level=1)
    doc.add_paragraph(
        "WidgetEnglish es una aplicación móvil para aprender inglés mediante microinteracciones: palabras, verbos, lotes de contenido, quizzes, tarjetas de estudio, estadísticas personales y un widget de pantalla de inicio. "
        "La app combina almacenamiento local con Room, configuración con DataStore, autenticación y sincronización con Firebase, y pronunciación mediante TextToSpeech."
    )
    add_callout(
        doc,
        "Idea central",
        "El usuario no necesita entrar siempre a una sesión larga de estudio: puede aprender, repasar y marcar progreso desde la app o desde el widget.",
    )

    doc.add_heading("2. Módulos principales", level=1)
    add_matrix_table(
        doc,
        ["Módulo", "Qué ofrece al usuario", "Datos o lógica asociada"],
        [
            ["Autenticación", "Registro, inicio de sesión, recuperación y cambio de contraseña.", "Firebase Authentication, perfiles locales y sincronización inicial de usuario."],
            ["Inicio", "Vista de bienvenida, resumen del progreso y acceso rápido a funciones.", "Lee usuario, lote activo, objetivo diario y datos de actividad."],
            ["Vocabulario", "Exploración de palabras y verbos con traducción, fonética y pronunciación.", "Room guarda palabras, verbos, lotes y progreso por contenido."],
            ["Lotes", "Selección de categorías o paquetes de aprendizaje.", "Un lote activo alimenta el widget y las sesiones de estudio."],
            ["Estudio y tarjetas", "Práctica en formato de tarjetas y sesiones configurables.", "Actualiza progreso, repasos, dominio y estadísticas de actividad."],
            ["Quiz", "Evaluación con preguntas, resultados y repaso de falladas.", "Registra correctas, incorrectas, quizzes completados y errores globales."],
            ["IA", "Ayuda conversacional o explicación asistida para practicar inglés.", "Módulo de cliente remoto de IA y pantallas de chat."],
            ["Perfil", "Ajustes de usuario, widget, objetivo diario y seguridad.", "DataStore para preferencias y Room/Firebase para datos de usuario."],
            ["Estadísticas personales", "Progreso, racha, precisión, lotes y actividad por semana/mes/año.", "Combina usuario, progreso, lotes y actividad diaria local."],
            ["Widget", "Aprendizaje desde la pantalla de inicio con palabra, fonética, traducción y acciones.", "AppWidgetProvider + RemoteViews + DataStore + selección de contenido."],
        ],
        [1700, 3650, 4010],
    )

    doc.add_heading("3. Flujo habitual de uso", level=1)
    add_numbered(
        doc,
        [
            "El usuario inicia sesión con correo o Google. La app crea o sincroniza su registro local y remoto.",
            "Elige o activa un lote de aprendizaje. Ese lote queda disponible para estudio dentro de la app y para el widget.",
            "Practica mediante vocabulario, tarjetas, quiz o widget. Cada interacción actualiza progreso local.",
            "Cuando se estudian elementos, se registra actividad diaria. Si se cumple la meta, la racha se actualiza.",
            "Las estadísticas relevantes se sincronizan con Firestore para mantener continuidad entre sesiones y alimentar vistas administrativas.",
        ],
    )

    doc.add_heading("4. Widget: modos visuales y acciones", level=1)
    doc.add_paragraph(
        "El widget muestra el contenido activo del usuario. Si no existe un lote activo, presenta un estado vacío con la indicación de activarlo desde la app. "
        "Cuando hay contenido, muestra palabra o verbo, fonética, traducción y progreso dentro de la sesión."
    )
    add_label_table(
        doc,
        [
            ["Tamaños", "Compacto, normal y grande. En modo automático el provider decide según ancho/alto del widget; también se puede forzar un tamaño desde preferencias."],
            ["Temas de color", "Azul, morado, verde, naranja, turquesa, rosa, índigo, rojo y variantes suaves/cristal/aurora/océano/oscuro."],
            ["Estilos visuales", "Clásico, minimalista, card suave, contraste alto y nocturno. Cada estilo ajusta fondo, texto, énfasis y contraste."],
            ["Visibilidad", "El usuario puede mostrar u ocultar nombre del lote, progreso, fonética y traducción."],
            ["Acciones", "Siguiente avanza al próximo contenido; sonido reproduce pronunciación; aprendida marca progreso; doble toque en la raíz abre la app."],
            ["Protección de audio", "El receptor TTS reutiliza un solo motor y descarta taps repetidos del mismo texto en una ventana corta para evitar sonidos duplicados."],
        ],
    )

    doc.add_heading("5. Modos de selección de contenido", level=1)
    add_matrix_table(
        doc,
        ["Modo", "Cómo elige contenido", "Uso recomendado"],
        [
            ["Secuencial", "Ordena por el campo de orden del lote. Cada día rota el punto de inicio según días transcurridos y objetivo. Prioriza pendientes y deja aprendidas como repaso.", "Usuarios que quieren avanzar de forma ordenada sin saltos."],
            ["Aleatorio", "Ordena el lote y luego lo mezcla con una semilla diaria. El resultado cambia por día, pero se mantiene estable durante ese día.", "Usuarios que prefieren variedad sin perder consistencia diaria."],
            ["Inteligente", "Calcula prioridad: nuevo/no visto tiene prioridad alta, difícil sube, en progreso queda al medio y aprendido baja. Ajusta por dominio, errores, pocos repasos y revisión vencida.", "Usuarios que quieren que la app empuje lo que más necesita práctica."],
        ],
        [1600, 5100, 2660],
    )

    doc.add_heading("6. Objetivo diario", level=1)
    doc.add_paragraph(
        "La app maneja un objetivo diario de estudio que sirve para medir actividad y racha. El objetivo efectivo depende de si el usuario está en modo automático o manual."
    )
    add_label_table(
        doc,
        [
            ["Automático", "Empieza en 5 y se mantiene entre 5 y 15. Si el día anterior hubo actividad pero no se cumplió la meta, baja. Si se cumple y la racha alcanza múltiplos de 3 días, sube."],
            ["Manual", "El usuario elige entre 10, 15, 20, 25, 30, 35 o 40. Si llega un valor no exacto, se normaliza a la opción más cercana."],
            ["Cumplimiento", "La actividad diaria suma elementos estudiados. Cuando elementosEstudiados >= objetivoDiario, se marca objetivoCumplido y se guarda fechaCumplimiento."],
            ["Racha", "Si se cumple el objetivo por primera vez en el día, la racha se incrementa si el último día cumplido fue ayer; si no, reinicia en 1."],
        ],
    )

    doc.add_heading("7. Datos que se guardan", level=1)
    add_bullets(
        doc,
        [
            "Usuario: nombre, correo, rol, estado activo, racha, palabras aprendidas, quizzes, lotes completados y porcentaje de progreso.",
            "Progreso por contenido: estado de aprendizaje, dominio, correctas, incorrectas, veces repasado, favorito, última y próxima revisión.",
            "Progreso por lote: si está activo o completado, porcentaje, contenidos aprendidos, total y fechas de estudio.",
            "Actividad diaria: elementos estudiados, tarjetas, preguntas de quiz, quizzes completados, objetivo diario y cumplimiento.",
            "Preferencias: lote activo, índice del widget, modo de selección, apariencia del widget y objetivo diario.",
        ],
    )

    doc.add_heading("8. Archivos fuente revisados", level=1)
    add_bullets(
        doc,
        [
            "README.md",
            "features/Screen.kt",
            "features/widget/WordWidgetProvider.kt y TtsReceiver.kt",
            "data/local/datastore/WidgetPreferences.kt, WidgetAppearancePreferences.kt, LearningPreferences.kt y DailyGoalPreferences.kt",
            "domain/learning/LearningContentSelector.kt",
            "data/repository/StreakRepository.kt",
            "data/local/entity/UsuarioEntity.kt, ProgresoUsuarioEntity.kt, ProgresoLoteEntity.kt y ActividadDiariaEntity.kt",
        ],
    )

    doc.save(USER_DOCX)


def build_admin_doc() -> None:
    doc = Document()
    style_document(doc, "WidgetEnglish - Estadísticas administrativas")
    add_title_block(
        doc,
        "WidgetEnglish: estadísticas administrativas",
        "Cómo se calculan, sincronizan y visualizan las métricas del panel admin",
        "Alcance: documento técnico-funcional para entender el pipeline de métricas que ve el administrador y las fuentes de datos que lo alimentan.",
    )

    doc.add_heading("1. Resumen del flujo de estadísticas", level=1)
    add_numbered(
        doc,
        [
            "El usuario estudia desde widget, tarjetas, vocabulario o quiz.",
            "La app actualiza Room: progreso de contenido, progreso de lote y actividad diaria.",
            "StreakRepository recalcula estadísticas de usuario, racha y cumplimiento de objetivo.",
            "Si existe data source remoto, se sincronizan actividad diaria y resumen de usuario en Firestore.",
            "El panel admin consulta usuarios, categorías globales y palabras con errores desde Firestore.",
            "AdminViewModel calcula totales, promedios, rankings y listas filtradas para las pantallas.",
        ],
    )

    doc.add_heading("2. Fuentes de datos", level=1)
    add_matrix_table(
        doc,
        ["Fuente", "Qué contiene", "Uso en admin"],
        [
            ["Firestore: usuarios", "Nombre, correo, activo, último acceso, palabras aprendidas, quizzes, lotes completados, racha y progreso.", "Base de totales, usuarios activos, rankings, actividad y cumplimiento."],
            ["Firestore: stats_globales/categorias/detalle", "Documento por lote/categoría con nombre, vecesEstudiada y última actualización.", "Pantalla de categorías más usadas."],
            ["Firestore: stats_globales/errores_palabras/detalle", "Documento por palabra con término, loteId, cantidadErrores y última actualización.", "Pantalla de palabras con más errores y filtro por categoría."],
            ["Firestore: usuarios/{uid}/actividadDiaria/{fecha}", "Actividad diaria sincronizada por usuario y fecha.", "Persistencia remota de actividad; útil para continuidad y auditoría."],
            ["Room local", "Usuario, progreso por contenido/lote y actividad diaria.", "Fuente primaria durante el uso de la app; se recalcula y sincroniza cuando hay actividad."],
        ],
        [2500, 4050, 2810],
    )

    doc.add_heading("3. Métricas calculadas para el resumen admin", level=1)
    add_matrix_table(
        doc,
        ["Métrica", "Cálculo actual", "Interpretación"],
        [
            ["Total usuarios", "Cantidad de documentos de usuarios no ADMIN.", "Tamaño de la base de usuarios gestionables."],
            ["Usuarios activos", "Conteo de usuarios con activo = true.", "Usuarios habilitados o marcados como activos."],
            ["Total palabras aprendidas", "Suma de palabrasAprendidas de todos los usuarios.", "Volumen acumulado de aprendizaje."],
            ["Total quizzes realizados", "Suma de quizzesRealizados.", "Participación en evaluaciones."],
            ["Total lotes completados", "Suma de lotesCompletados.", "Avance acumulado por paquetes de contenido."],
            ["Promedio palabras/usuario", "totalPalabras / totalUsuarios, entero.", "Promedio simple de aprendizaje por usuario."],
            ["Promedio quizzes/usuario", "totalQuizzes / totalUsuarios, entero.", "Promedio simple de evaluaciones por usuario."],
            ["Porcentaje usuarios activos", "usuariosActivos / totalUsuarios * 100.", "Proporción de usuarios habilitados/activos."],
            ["Cumplimiento promedio", "Promedio de porcentajeProgreso de usuarios.", "Avance promedio global declarado por usuario."],
        ],
        [2300, 3950, 3110],
    )

    doc.add_heading("4. Ranking y actividad", level=1)
    add_label_table(
        doc,
        [
            ["Ranking", "Ordena el top 10 por palabras aprendidas, quizzes realizados o racha actual. La pantalla permite cambiar el criterio con chips."],
            ["Actividad", "Ordena el top 10 por último acceso, racha actual o porcentaje de progreso. Sirve para revisar participación reciente o cumplimiento."],
            ["Resumen", "El dashboard muestra usuarios, aprendidas, quizzes, activos, lotes y tamaño del ranking, además de una vista previa de los 3 usuarios más activos."],
            ["Actualización manual", "Las pantallas de ranking, actividad, categorías y errores incluyen acción de refrescar con forzarActualizacion = true."],
        ],
    )

    doc.add_heading("5. Categorías y errores globales", level=1)
    add_label_table(
        doc,
        [
            ["Categorías más usadas", "Se leen ordenadas por vecesEstudiada descendente. La UI muestra ranking, número de sesiones y una barra proporcional al máximo de la lista."],
            ["Registro de categoría", "StreakRepository llama incrementarUsoCategoria(loteId, nombre), que incrementa vecesEstudiada en stats_globales/categorias/detalle/{loteId}."],
            ["Palabras con más errores", "Se leen hasta 50 palabras ordenadas por cantidadErrores descendente. La UI permite filtrar por loteId/categoría."],
            ["Registro de error", "En quiz, al fallar una palabra se invoca registrarErrorPalabraGlobal, que incrementa cantidadErrores en stats_globales/errores_palabras/detalle/{palabraId}."],
        ],
    )

    doc.add_heading("6. Actividad diaria, racha y objetivo", level=1)
    doc.add_paragraph(
        "La actividad diaria se guarda por usuario y fecha. Cada registro acumula elementos estudiados, tarjetas, preguntas respondidas, quizzes completados y el estado del objetivo diario."
    )
    add_matrix_table(
        doc,
        ["Elemento", "Lógica"],
        [
            ["Actividad del día", "Si no existe registro para hoy, se crea. Si existe, se suman elementos, tarjetas, preguntas y quizzes."],
            ["Objetivo efectivo", "Se resuelve desde DailyGoalPreferences si no se pasa un objetivo específico. Puede ser automático o manual."],
            ["Cumplimiento", "objetivoCumplido pasa a true cuando elementosEstudiados acumulados alcanza objetivoDiario."],
            ["Racha", "Al cumplir por primera vez en el día, si la última fecha de racha fue ayer, suma 1; si no, reinicia en 1."],
            ["Ajuste automático", "Si ayer hubo actividad y no se cumplió el objetivo, baja. Si ayer se cumplió y la racha actual es múltiplo de 3, sube."],
            ["Sincronización", "Se sube actividadDiaria/{fecha} y luego el resumen del usuario con racha, palabras, quizzes, lotes y progreso."],
        ],
        [2300, 7060],
    )

    doc.add_heading("7. Estadísticas personales como contexto", level=1)
    add_bullets(
        doc,
        [
            "La pantalla personal calcula palabras aprendidas, quizzes, progreso, racha, lotes completados y precisión global.",
            "La precisión se calcula como correctas / (correctas + incorrectas) * 100.",
            "El estado de aprendizaje se distribuye en aprendidas, en progreso, difíciles y no vistas.",
            "El historial se agrupa por semana, mes o año usando actividad diaria: días de la semana, semanas del mes o meses del año.",
            "El progreso por lotes se ordena por porcentaje, contenidos aprendidos y estado activo.",
        ],
    )

    doc.add_heading("8. Consideraciones y límites actuales", level=1)
    add_bullets(
        doc,
        [
            "AdminFirestoreDataSource excluye usuarios con rol ADMIN para que el panel se centre en usuarios de aprendizaje.",
            "AdminViewModel evita cargas duplicadas con cargaEnCurso y cachea datos salvo que se fuerce actualización.",
            "Si fallan categorías o errores globales, el panel conserva los datos de usuarios y muestra un error parcial.",
            "La pantalla de errores limita la consulta a 50 términos con más errores.",
            "El cumplimiento promedio depende del campo porcentajeProgreso sincronizado por usuario, no de una consulta histórica completa.",
            "Los totales administrativos dependen de la última sincronización de cada usuario; si un usuario no ha sincronizado, el admin puede ver datos atrasados.",
        ],
    )

    doc.add_heading("9. Archivos fuente revisados", level=1)
    add_bullets(
        doc,
        [
            "features/admin/AdminViewModel.kt y AdminUiState.kt",
            "features/admin/AdminDashboardScreen.kt, ranking/AdminRankingScreen.kt, activity/AdminActivityScreen.kt",
            "features/admin/stats/AdminCategoriasScreen.kt y AdminErroresScreen.kt",
            "data/remote/firestore/AdminFirestoreDataSource.kt y EstadisticasFirestoreDataSource.kt",
            "data/repository/StreakRepository.kt",
            "features/profile/statistics/viewmodel/StatisticsViewModel.kt y model/StatisticsUiState.kt",
            "data/local/entity/ActividadDiariaEntity.kt, UsuarioEntity.kt, ProgresoUsuarioEntity.kt y ProgresoLoteEntity.kt",
        ],
    )

    doc.save(ADMIN_DOCX)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_user_doc()
    build_admin_doc()
    print(USER_DOCX)
    print(ADMIN_DOCX)


if __name__ == "__main__":
    main()
